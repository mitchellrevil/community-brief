import os
import asyncio
import inspect
from typing import Optional, AsyncGenerator
from azure.storage.blob.aio import BlobServiceClient, BlobClient
from azure.storage.blob import BlobSasPermissions, ContentSettings, generate_blob_sas
from azure.identity.aio import DefaultAzureCredential
from azure.core.exceptions import AzureError
from datetime import UTC, datetime, timedelta
from urllib.parse import urlparse
from azure.core.exceptions import ResourceNotFoundError
from ...core.config import AppConfig
from ...core.logging import get_logger
from ...utils.cache_utils import TTLCache

BLOB_SERVICE_ERRORS = (AzureError, RuntimeError, OSError, ValueError, TypeError)
BLOB_STREAM_ERRORS = (AzureError, RuntimeError, OSError, ValueError, TypeError)
MARKDOWN_PARSE_ERRORS = (ValueError, TypeError, AttributeError)
MARKDOWN_IMPORT_ERRORS = (ImportError, ModuleNotFoundError)

_sas_url_cache = TTLCache[str](default_ttl=300.0)


class StorageService:
    def __init__(self, config: AppConfig):
        self.config = config
        self.logger = get_logger(__name__)
        
        if config.azure_storage_key:
            self.logger.info("azure_storage_key_auth_configured")
            self.credential = config.azure_storage_key
        else:
            self.logger.info("azure_storage_managed_identity_auth_configured")
            self.credential = DefaultAzureCredential()

        self.blob_service_client = BlobServiceClient(
            account_url=self.config.azure_storage_account_url, credential=self.credential
        )

        self._user_delegation_key = None
        self._user_delegation_key_expiration: Optional[datetime] = None
        self._delegation_key_lock = asyncio.Lock()

    async def close(self):
        """Close the async blob service client."""
        await self.blob_service_client.close()
        # Close credential if it supports an async close() method.
        # Avoid isinstance checks which break when tests patch the class with mocks.
        try:
            close_fn = getattr(self.credential, "close", None)
            if close_fn and callable(close_fn):
                await close_fn()
        except BLOB_SERVICE_ERRORS:
            # Best-effort close, ignore errors during teardown
            pass

    async def generate_sas_token(self, blob_url: str) -> Optional[str]:
        """Generate SAS token for a blob URL using managed identity"""
        try:
            if not blob_url:
                return None

            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")
            if len(path_parts) < 2:
                self.logger.warning("blob_url_invalid", blob_url=blob_url)
                return None

            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])
            account_name = parsed_url.netloc.split(".")[0]

            now = datetime.now(UTC)
            start_time = now - timedelta(minutes=5)
            expiry_time = now + timedelta(hours=8)

            if isinstance(self.credential, str):
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_name,
                    account_key=self.credential,
                    permission=BlobSasPermissions(read=True),
                    start=start_time,
                    expiry=expiry_time,
                )
            else:
                user_delegation_key = await self._get_user_delegation_key(start_time)
                if user_delegation_key is None:
                    return None

                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_name,
                    user_delegation_key=user_delegation_key,
                    permission=BlobSasPermissions(read=True),
                    start=start_time,
                    expiry=expiry_time,
                )

            self.logger.debug(
                "blob_sas_token_generated",
                container=container_name,
                blob_name=blob_name[:50],
                start_time=start_time.isoformat(),
                expiry_time=expiry_time.isoformat(),
            )
            return sas_token

        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_sas_token_generation_failed",
                blob_url=blob_url,
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            return None

    @staticmethod
    def _sanitize_blob_segment(value: str) -> str:
        return "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in (value or ""))

    def _sanitize_storage_filename(self, original_filename: str) -> str:
        return (original_filename or "upload.bin").replace(" ", "_")

    def _build_direct_upload_blob_name(self, original_filename: str, owner_user_id: str) -> str:
        sanitized_filename = self._sanitize_storage_filename(original_filename)
        owner_segment = self._sanitize_blob_segment(owner_user_id or "anonymous") or "anonymous"
        current_date = datetime.now(UTC).strftime("%Y-%m-%d")
        timestamp = datetime.now(UTC).strftime("%H%M%S_%f")[:-3]
        file_name_without_ext = os.path.splitext(sanitized_filename)[0]
        return (
            f"direct/{owner_segment}/{current_date}/"
            f"{file_name_without_ext}_{timestamp}/{sanitized_filename}"
        )

    async def generate_upload_sas(self, original_filename: str, owner_user_id: str) -> dict:
        """Generate a write-only SAS URL for direct client-to-blob upload.

        Returns a dict with ``blob_url``, ``sas_url`` (full URL with token),
        ``blob_name``, and ``container`` so the caller can upload directly from
        the browser and later reference the blob.
        """
        container_name = self.config.azure_storage_recordings_container
        blob_name = self._build_direct_upload_blob_name(original_filename, owner_user_id)

        parsed = urlparse(self.config.azure_storage_account_url)
        account_name = parsed.netloc.split(".")[0]

        now = datetime.now(UTC)
        start_time = now - timedelta(minutes=5)
        expiry_time = now + timedelta(hours=1)  # short-lived for security

        permissions = BlobSasPermissions(create=True, write=True)

        if isinstance(self.credential, str):
            sas_token = generate_blob_sas(
                account_name=account_name,
                container_name=container_name,
                blob_name=blob_name,
                account_key=self.credential,
                permission=permissions,
                start=start_time,
                expiry=expiry_time,
            )
        else:
            udk = await self._get_user_delegation_key(start_time)
            if udk is None:
                raise RuntimeError("Unable to obtain user delegation key for SAS generation")
            sas_token = generate_blob_sas(
                account_name=account_name,
                container_name=container_name,
                blob_name=blob_name,
                user_delegation_key=udk,
                permission=permissions,
                start=start_time,
                expiry=expiry_time,
            )

        blob_url = f"{self.config.azure_storage_account_url.rstrip('/')}/{container_name}/{blob_name}"
        sas_url = f"{blob_url}?{sas_token}"

        self.logger.info(
            "blob_upload_sas_token_generated",
            blob_name=blob_name,
            expiry=expiry_time.isoformat(),
        )
        return {
            "blob_url": blob_url,
            "sas_url": sas_url,
            "blob_name": blob_name,
            "container": container_name,
            "expiry": expiry_time.isoformat(),
        }

    def is_expected_direct_upload_blob(
        self,
        *,
        blob_url: str,
        original_filename: str,
        owner_user_id: str,
    ) -> bool:
        if not blob_url:
            return False

        parsed_url = urlparse(blob_url)
        expected_account = urlparse(self.config.azure_storage_account_url).netloc
        path_parts = parsed_url.path.strip("/").split("/")
        if parsed_url.netloc != expected_account or len(path_parts) < 5:
            return False

        container_name = path_parts[0]
        blob_name = "/".join(path_parts[1:])
        expected_prefix = f"direct/{self._sanitize_blob_segment(owner_user_id or 'anonymous')}/"
        expected_filename = self._sanitize_storage_filename(original_filename)

        return (
            container_name == self.config.azure_storage_recordings_container
            and blob_name.startswith(expected_prefix)
            and blob_name.endswith(f"/{expected_filename}")
        )

    async def verify_blob_exists(self, blob_url: str) -> int:
        """Verify a blob exists and return its size in bytes.

        Raises ``FileNotFoundError`` if the blob does not exist.
        """
        parsed_url = urlparse(blob_url)
        path_parts = parsed_url.path.strip("/").split("/")
        if len(path_parts) < 2:
            raise ValueError(f"Invalid blob URL: {blob_url}")

        container_name = path_parts[0]
        blob_name = "/".join(path_parts[1:])

        container_client = self.blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(blob_name)
        try:
            props = await blob_client.get_blob_properties()
            return props.size
        except ResourceNotFoundError:
            raise FileNotFoundError(f"Blob not found: {blob_url}")

    async def set_blob_metadata(self, blob_url: str, metadata: dict) -> bool:
        """Set custom metadata properties on a blob.

        This is used to attach correlation_id and file_name to blobs so Azure Functions
        can retrieve jobs by blob metadata instead of relying on URL matching.

        Args:
            blob_url: Full URL to the blob (e.g., https://<storage>.blob.core.windows.net/<container>/<blob>)
            metadata: Dictionary of metadata key-value pairs. Keys must be lowercase ASCII letters/numbers.

        Returns:
            True if metadata was set successfully, False on any error.

        Note:
            This method is non-fatal: failures are logged but don't raise exceptions.
            The upload flow should succeed even if metadata attachment fails.
        """
        if not blob_url:
            self.logger.warning("blob_metadata_set_empty_url")
            return False

        try:
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")

            if len(path_parts) < 2:
                self.logger.warning("blob_metadata_url_invalid", blob_url=blob_url)
                return False

            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])

            container_client = self.blob_service_client.get_container_client(container_name)
            blob_client = container_client.get_blob_client(blob_name)

            await blob_client.set_blob_metadata(metadata)

            self.logger.info(
                "blob_metadata_set_succeeded",
                blob_url=blob_url[:80],
                metadata_keys=list(metadata.keys()),
                correlation_id=metadata.get("correlationid", "unknown"),
                blob_filename=metadata.get("filename", "unknown"),
            )
            return True

        except ResourceNotFoundError:
            self.logger.warning(
                "blob_metadata_set_blob_not_found",
                blob_url=blob_url[:80],
                correlation_id=metadata.get("correlationid", "unknown") if metadata else "unknown",
            )
            return False
        except AzureError as e:
            self.logger.warning(
                "blob_metadata_set_azure_error",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                correlation_id=metadata.get("correlationid", "unknown") if metadata else "unknown",
            )
            return False
        except BLOB_SERVICE_ERRORS as e:
            self.logger.warning(
                "blob_metadata_set_unexpected_error",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                correlation_id=metadata.get("correlationid", "unknown") if metadata else "unknown",
                exc_info=True
            )
            return False

    async def add_sas_token_to_url(self, blob_url: str) -> str:
        """Add SAS token to blob URL if not already present"""
        if not blob_url:
            return blob_url

        if '?' in blob_url:
            self.logger.debug("blob_url_already_has_query", blob_url=blob_url[:100])
            return blob_url

        cached = await _sas_url_cache.get(blob_url)
        if cached is not None:
            return cached

        sas_token = await self.generate_sas_token(blob_url)
        if sas_token:
            sas_url = f"{blob_url}?{sas_token}"
            await _sas_url_cache.set(blob_url, sas_url)
            return sas_url
        
        self.logger.warning("blob_sas_token_unavailable", blob_url=blob_url)
        return blob_url

    async def _get_user_delegation_key(self, start_time: datetime):
        """Return a cached user delegation key or fetch a new one if needed."""
        if isinstance(self.credential, str):
            return None

        try:
            async with self._delegation_key_lock:
                now = datetime.now(UTC)
                cached_expiration = self._user_delegation_key_expiration
                if cached_expiration is not None and cached_expiration.tzinfo is None:
                    cached_expiration = cached_expiration.replace(tzinfo=UTC)
                if (
                    self._user_delegation_key is not None
                    and cached_expiration is not None
                    and now < cached_expiration - timedelta(minutes=1)
                ):
                    return self._user_delegation_key

                key_expiry_time = now + timedelta(hours=12)
                self._user_delegation_key = await self.blob_service_client.get_user_delegation_key(
                    key_start_time=start_time,
                    key_expiry_time=key_expiry_time,
                )
                self._user_delegation_key_expiration = key_expiry_time
                return self._user_delegation_key
        except BLOB_SERVICE_ERRORS as exc:
            self.logger.error("user_delegation_key_refresh_failed", exc_info=True, error=str(exc))
            return None

    async def download_blob_bytes(self, blob_url: str) -> bytes:
        """Download raw bytes from a blob URL.
        
        Args:
            blob_url: Full URL to the blob to download
            
        Returns:
            Raw bytes content of the blob
            
        Raises:
            FileNotFoundError: If blob does not exist
            ValueError: If blob URL is invalid
            Exception: For other download errors
        """
        if not blob_url:
            raise ValueError("Blob URL cannot be empty")
            
        try:
            self.logger.debug("blob_bytes_download_started", blob_url=blob_url[:80])
            
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")
            
            if len(path_parts) < 2:
                self.logger.error("blob_url_invalid", blob_url=blob_url)
                raise ValueError(f"Invalid blob URL format: {blob_url}")
                
            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])
            
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            blob_client = container_client.get_blob_client(blob_name)
            
            stream = await blob_client.download_blob()
            blob_data = await stream.readall()
            
            self.logger.info(
                "blob_bytes_download_succeeded",
                blob_url=blob_url[:80],
                size_bytes=len(blob_data),
            )
            
            return blob_data
            
        except ResourceNotFoundError as e:
            self.logger.error("blob_bytes_download_not_found", blob_url=blob_url, exc_info=True)
            raise FileNotFoundError(f"Blob not found: {blob_url}") from e
        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_bytes_download_failed",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True
            )
            raise

    async def upload_blob_bytes(self, original_filename: str, file_data: bytes) -> str:
        """Upload raw bytes to blob storage with automatic naming.
        
        Args:
            original_filename: Original filename for the blob
            file_data: Raw bytes to upload
            
        Returns:
            Full blob URL of the uploaded blob
            
        Raises:
            Exception: If upload fails
        """
        if not original_filename:
            raise ValueError("Original filename cannot be empty")
        if not isinstance(file_data, bytes):
            raise ValueError("File data must be bytes")
            
        try:
            container_name = self.config.azure_storage_recordings_container
            
            # Generate blob name with date and nested structure including timestamp for uniqueness
            sanitized_filename = original_filename.replace(" ", "_")
            current_date = datetime.now(UTC).strftime("%Y-%m-%d")
            timestamp = datetime.now(UTC).strftime("%H%M%S_%f")[:-3]  # HHMMSS_milliseconds
            file_name_without_ext = os.path.splitext(sanitized_filename)[0]
            blob_name = f"{current_date}/{file_name_without_ext}_{timestamp}/{sanitized_filename}"

            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            blob_client = container_client.get_blob_client(blob_name)

            self.logger.info(
                "blob_bytes_upload_started",
                blob_name=blob_name,
                size_bytes=len(file_data),
            )
            
            await blob_client.upload_blob(file_data, overwrite=True)
            
            blob_url = blob_client.url
            self.logger.info(
                "blob_bytes_upload_succeeded",
                blob_url=blob_url[:80],
                size_bytes=len(file_data),
            )
            
            return blob_url

        except AzureError as e:
            self.logger.error(
                "blob_bytes_upload_azure_error",
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            raise
        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_bytes_upload_failed",
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            raise

    async def download_text_from_blob(self, blob_url: str) -> Optional[str]:
        """Download text content from a blob URL."""
        if not blob_url:
            return None
            
        try:
            self.logger.debug("blob_text_download_started", blob_url=blob_url[:80])
            
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")
            
            if len(path_parts) < 2:
                self.logger.error("blob_url_invalid", blob_url=blob_url)
                return None
                
            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])
            
            # Get container client (some mocks return coroutine, handle awaitables)
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            # In Azure SDK, get_container_client is synchronous but returns an async client.
            # However, some mocks might make it async. We assume standard SDK behavior here:
            # it returns a client immediately. If it's an async client, we use it.
            # If tests mock it as a coroutine, they should be fixed.
            
            blob_client = container_client.get_blob_client(blob_name)
            
            stream = await blob_client.download_blob()
            blob_data = await stream.readall()
            text_content = blob_data.decode('utf-8')
            
            self.logger.info(
                "blob_text_download_succeeded",
                blob_url=blob_url[:80],
                content_length=len(text_content),
            )
            
            return text_content
            
        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_text_download_failed",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True
            )
            return None

    async def upload_text_to_blob(
        self,
        blob_url: str,
        text_content: str,
        *,
        content_type: str = "text/plain; charset=utf-8",
    ) -> str:
        """Replace an existing blob with UTF-8 text content."""
        if not blob_url:
            raise ValueError("Blob URL cannot be empty")

        try:
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")

            if len(path_parts) < 2:
                self.logger.error("blob_url_invalid", blob_url=blob_url)
                raise ValueError(f"Invalid blob URL format: {blob_url}")

            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])

            container_client = self.blob_service_client.get_container_client(container_name)
            blob_client = container_client.get_blob_client(blob_name)
            payload = text_content.encode("utf-8")

            await blob_client.upload_blob(
                payload,
                overwrite=True,
                content_settings=ContentSettings(content_type=content_type),
            )

            self.logger.info(
                "blob_text_upload_succeeded",
                blob_url=blob_url[:80],
                size_bytes=len(payload),
            )
            return blob_client.url

        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_text_upload_failed",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            raise

    async def download_docx_text_from_blob(self, blob_url: str) -> Optional[str]:
        """Download .docx file from blob and extract text content.
        
        Args:
            blob_url: Full URL to the .docx blob to download
            
        Returns:
            Extracted text content from the .docx, or None if download/extraction fails
        """
        if not blob_url:
            return None
            
        try:
            self.logger.debug("blob_docx_download_started", blob_url=blob_url[:80])
            
            # Extract container and blob name from URL path
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")
            
            if len(path_parts) < 2:
                self.logger.error("blob_url_invalid", blob_url=blob_url)
                return None
                
            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])
            
            # Get blob client
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            
            blob_client = container_client.get_blob_client(blob_name)
            
            # Download .docx file as bytes
            stream = await blob_client.download_blob()
            blob_data = await stream.readall()
            
            # Extract text from .docx (run in thread since python-docx is sync)
            def _extract_text(data):
                from docx import Document
                import io
                doc = Document(io.BytesIO(data))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                return "\n\n".join(text_parts)

            extracted_text = await asyncio.to_thread(_extract_text, blob_data)
            
            self.logger.info(
                "blob_docx_text_extract_succeeded",
                blob_url=blob_url[:80],
                content_length=len(extracted_text),
            )
            
            return extracted_text
            
        except ImportError as ie:
            self.logger.error(
                "blob_docx_dependency_missing",
                blob_url=blob_url[:80],
                error=str(ie),
            )
            return None
        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_docx_text_extract_failed",
                blob_url=blob_url[:80],
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True
            )
            return None

    async def upload_file(self, file_path: str, original_filename: str) -> str:
        """Upload a file to blob storage"""
        try:
            container_client = self.blob_service_client.get_container_client(
                self.config.azure_storage_recordings_container
            )            # Sanitize filename - replace spaces with underscores
            sanitized_filename = original_filename.replace(" ", "_")
            self.logger.debug(
                "blob_upload_filename_sanitized",
                original_filename=original_filename,
                sanitized_filename=sanitized_filename,
            )
            
            # Generate blob name with date and nested structure including timestamp for uniqueness
            current_date = datetime.now(UTC).strftime("%Y-%m-%d")
            timestamp = datetime.now(UTC).strftime("%H%M%S_%f")[:-3]  # HHMMSS_milliseconds
            file_name_without_ext = os.path.splitext(sanitized_filename)[0]
            blob_name = f"{current_date}/{file_name_without_ext}_{timestamp}/{sanitized_filename}"

            blob_client = container_client.get_blob_client(blob_name)

            self.logger.info("blob_file_upload_started", blob_name=blob_name)
            
            # Read file in thread to avoid blocking (max 500MB per config)
            async with asyncio.Lock():
                def _read_file():
                    with open(file_path, "rb") as data:
                        return data.read()
                
                file_data = await asyncio.to_thread(_read_file)
                await blob_client.upload_blob(file_data, overwrite=True)

            return blob_client.url

        except AzureError as e:
            self.logger.error(
                "blob_file_upload_azure_error",
                error=str(e),
                error_type=type(e).__name__,
            )
            raise
        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_file_upload_failed",
                error=str(e),
                error_type=type(e).__name__,
            )
            raise

    async def generate_and_upload_docx(self, analysis_text: str, blob_name: str, add_title: bool = True) -> str:
        """Generate a DOCX from analysis text and upload to blob storage. Return the blob URL.
        
        Args:
            analysis_text: The text content to convert to DOCX
            blob_name: The name/path for the blob in storage
            add_title: Whether to add "Analysis Report" title at the top (True for new, False for edited)
        """
        try:
            def _generate_docx():
                from docx import Document
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                import importlib
                import io

                doc = Document()
                if add_title:
                    title = doc.add_heading('Analysis Report', 0)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                try:
                    from markdown_it import MarkdownIt

                    try:
                        md = MarkdownIt("commonmark", {"linkify": True})
                    except MARKDOWN_IMPORT_ERRORS:
                        md = MarkdownIt({"linkify": True})

                    for rule in ("table", "strikethrough"):
                        try:
                            md.enable(rule)
                        except MARKDOWN_PARSE_ERRORS:
                            self.logger.debug("blob_docx_markdown_rule_unavailable", rule=rule)

                    tasklist_loaded = False
                    tasklist_candidates = [
                        ("mdit_py_plugins.tasklist", "tasklist_plugin"),
                        ("mdit_py_plugins.tasklists", "plugin"),
                        ("mdit_py_plugins.tasklists", "tasklist_plugin"),
                    ]
                    for mod_name, attr in tasklist_candidates:
                        try:
                            mod = importlib.import_module(mod_name)
                            plugin = getattr(mod, attr)
                            md.use(plugin)
                            tasklist_loaded = True
                            self.logger.info(
                                "blob_docx_tasklist_plugin_loaded",
                                module=mod_name,
                                attribute=attr,
                            )
                            break
                        except MARKDOWN_IMPORT_ERRORS:
                            continue
                        except MARKDOWN_PARSE_ERRORS:
                            continue

                    if not tasklist_loaded:
                        self.logger.debug("blob_docx_tasklist_plugin_unavailable")

                    tokens = md.parse(analysis_text)
                except ModuleNotFoundError:
                    self.logger.warning("blob_docx_markdown_dependency_missing")
                    paragraphs = [p for p in analysis_text.split("\n\n") if p.strip()]
                    for para in paragraphs:
                        p = doc.add_paragraph()
                        fake_inline = type(
                            "InlineToken",
                            (),
                            {"children": [type("Child", (), {"type": "text", "content": para})]},
                        )
                        self._render_inline_tokens(p, fake_inline)

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    return buffer.getvalue()
                except MARKDOWN_PARSE_ERRORS as parse_err:
                    self.logger.error("blob_docx_markdown_parse_failed", error=str(parse_err))
                    raise ValueError(f"Markdown parsing error: {parse_err}") from parse_err

                list_stack: list[dict] = []

                i = 0
                while i < len(tokens):
                    tok = tokens[i]

                    if tok.type == 'heading_open':
                        level = 1
                        try:
                            tag = tok.tag
                            if tag and tag.startswith('h'):
                                level = max(1, min(3, int(tag[1:])))
                        except MARKDOWN_PARSE_ERRORS:
                            level = 1

                        text = ''
                        if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                            text = self._collect_plain_text(tokens[i + 1])

                        doc.add_heading(text.strip(), level=level)
                        i += 3
                        continue

                    if tok.type == 'paragraph_open':
                        p = doc.add_paragraph()
                        if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                            self._render_inline_tokens(p, tokens[i + 1])
                        i += 3
                        continue

                    if tok.type in ('ordered_list_open', 'bullet_list_open'):
                        list_type = 'ordered' if tok.type == 'ordered_list_open' else 'bullet'
                        level = len(list_stack) + 1
                        list_stack.append({'type': list_type, 'level': level})
                        i += 1
                        continue

                    if tok.type == 'list_item_open':
                        current = list_stack[-1] if list_stack else {'type': 'bullet', 'level': 1}
                        base_style = 'List Number' if current['type'] == 'ordered' else 'List Bullet'
                        style = base_style if current['level'] == 1 else f"{base_style} {current['level']}"

                        task_state = None
                        if isinstance(getattr(tok, "meta", None), dict) and "checked" in tok.meta:
                            task_state = bool(tok.meta.get("checked"))

                        try:
                            doc_style_names = {s.name for s in doc.styles}
                        except MARKDOWN_PARSE_ERRORS:
                            doc_style_names = set()

                        if style in doc_style_names:
                            use_style = style
                        elif base_style in doc_style_names:
                            use_style = base_style
                        else:
                            use_style = None

                        if use_style:
                            try:
                                p = doc.add_paragraph(style=use_style)
                            except MARKDOWN_PARSE_ERRORS:
                                p = doc.add_paragraph()
                        else:
                            p = doc.add_paragraph()

                        if task_state is not None:
                            prefix = "[x] " if task_state else "[ ] "
                            p.add_run(prefix)
                            setattr(p, "_task_prefix_added", True)

                        if i + 1 < len(tokens) and tokens[i + 1].type == 'paragraph_open':
                            if i + 2 < len(tokens) and tokens[i + 2].type == 'inline':
                                self._render_inline_tokens(p, tokens[i + 2])
                            i += 4
                        elif i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                            self._render_inline_tokens(p, tokens[i + 1])
                            i += 2
                        else:
                            i += 1
                        continue

                    if tok.type in ('ordered_list_close', 'bullet_list_close'):
                        if list_stack:
                            list_stack.pop()
                        i += 1
                        continue

                    if tok.type in ('fence', 'code_block'):
                        p = doc.add_paragraph()
                        self._render_code_block(p, tok.content or '')
                        i += 1
                        continue

                    if tok.type == 'table_open':
                        table_data = self._parse_table_tokens(tokens, i)
                        if table_data:
                            self._render_table(doc, table_data)
                            i = table_data['end_index']
                        else:
                            i += 1
                        continue

                    i += 1

                buffer = io.BytesIO()
                doc.save(buffer)
                return buffer.getvalue()

            docx_content = await asyncio.to_thread(_generate_docx)

            # Upload DOCX
            container_client = self.blob_service_client.get_container_client(
                self.config.azure_storage_recordings_container
            )

            blob_client = container_client.get_blob_client(blob_name)

            self.logger.info("blob_docx_upload_started", blob_name=blob_name)
            await blob_client.upload_blob(docx_content, overwrite=True)
            return blob_client.url

        except BLOB_SERVICE_ERRORS as e:
            self.logger.error(
                "blob_docx_upload_failed",
                blob_name=blob_name,
                error=str(e),
                error_type=type(e).__name__,
            )
            raise

    def _collect_plain_text(self, inline_token) -> str:
        """Collect plain text from an inline token."""
        text_parts = []
        for child in getattr(inline_token, 'children', []) or []:
            if child.type in ('text', 'code_inline'):
                text_parts.append(child.content)
            elif child.type in ('softbreak', 'hardbreak'):
                text_parts.append('\n')
            elif child.type == 'link_open':
                continue
        return ''.join(text_parts)

    def _render_inline_tokens(self, paragraph, inline_token) -> None:
        """Render inline markdown tokens into a python-docx paragraph."""
        from docx.shared import Pt, RGBColor

        children = getattr(inline_token, 'children', []) or []
        i = 0
        while i < len(children):
            child = children[i]

            if child.type == 'text':
                paragraph.add_run(child.content)
                i += 1
            elif child.type == 'code_inline':
                run = paragraph.add_run(child.content)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
            elif child.type == 'strong_open':
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'strong_close')
                run = paragraph.add_run(text)
                run.bold = True
                i += tokens_consumed
            elif child.type == 'em_open':
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'em_close')
                run = paragraph.add_run(text)
                run.italic = True
                i += tokens_consumed
            elif child.type in ('softbreak', 'hardbreak'):
                paragraph.add_run('\n')
                i += 1
            elif child.type == 'link_open':
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'link_close')
                href = self._get_token_attr(child, "href")
                if href:
                    self._add_hyperlink(paragraph, href, text or href)
                else:
                    paragraph.add_run(text)
                i += tokens_consumed
            elif child.type == 's_open':
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 's_close')
                run = paragraph.add_run(text)
                run.font.strike = True
                i += tokens_consumed
            elif child.type == 'task_list_item_checkbox':
                checked = False
                if isinstance(getattr(child, "meta", None), dict):
                    checked = bool(child.meta.get("checked"))
                if not getattr(paragraph, "_task_prefix_added", False):
                    prefix = "[x] " if checked else "[ ] "
                    paragraph.add_run(prefix)
                    setattr(paragraph, "_task_prefix_added", True)
                i += 1
            else:
                i += 1

    def _get_token_attr(self, token, attr_name: str) -> Optional[str]:
        """Return a token attribute value if present."""
        if hasattr(token, "attrGet"):
            try:
                value = token.attrGet(attr_name)
                if value:
                    return value
            except MARKDOWN_PARSE_ERRORS:
                pass

        attrs = getattr(token, "attrs", None) or []
        for key, value in attrs:
            if key == attr_name:
                return value
        return None

    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        """Add a hyperlink to a paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        new_run.append(r_pr)

        new_text = OxmlElement("w:t")
        new_text.text = text
        new_run.append(new_text)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _gather_text_between_with_count(self, children: list, start_index: int, end_type: str) -> tuple[str, int]:
        """Gather plain text between matching inline token delimiters."""
        text_parts = []
        open_count = 1
        start_token_type = children[start_index].type
        tokens_consumed = 1

        for i in range(start_index + 1, len(children)):
            token = children[i]
            tokens_consumed += 1

            if token.type == start_token_type:
                open_count += 1
            elif token.type == end_type:
                open_count -= 1
                if open_count == 0:
                    break
            elif token.type in ('text', 'code_inline'):
                text_parts.append(token.content)

        return ''.join(text_parts), tokens_consumed

    def _render_code_block(self, paragraph, content: str) -> None:
        """Render fenced/code block content in monospace formatting."""
        from docx.shared import Pt

        run = paragraph.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    def _parse_table_tokens(self, tokens: list, start_index: int) -> Optional[dict]:
        """Parse markdown-it table tokens into rows for python-docx."""
        rows = []
        current_row = []
        in_header = True
        i = start_index

        while i < len(tokens):
            tok = tokens[i]

            if tok.type == 'table_open':
                i += 1
                continue

            if tok.type == 'thead_open':
                in_header = True
                i += 1
                continue

            if tok.type == 'tbody_open':
                in_header = False
                i += 1
                continue

            if tok.type == 'tr_open':
                current_row = []
                i += 1
                continue

            if tok.type in ('th_open', 'td_open'):
                cell_content = ''
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    cell_content = self._collect_plain_text(tokens[i + 1])
                current_row.append(cell_content)
                i += 3
                continue

            if tok.type == 'tr_close':
                if current_row:
                    rows.append({
                        'cells': current_row,
                        'is_header': in_header,
                    })
                i += 1
                continue

            if tok.type in ('thead_close', 'tbody_close'):
                i += 1
                continue

            if tok.type == 'table_close':
                return {
                    'rows': rows,
                    'end_index': i + 1,
                }

            i += 1

        return None

    def _render_table(self, doc, table_data: dict) -> None:
        """Render a table in the Word document."""
        from docx.shared import Inches

        rows = table_data['rows']
        if not rows:
            return

        num_rows = len(rows)
        num_cols = max(len(row['cells']) for row in rows) if rows else 0
        if num_rows == 0 or num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        for col in table.columns:
            col.width = Inches(1.0)

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_content in enumerate(row_data['cells']):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_content

                    if row_data['is_header']:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    async def stream_blob_content(
        self, file_blob_url: str
    ) -> AsyncGenerator[bytes, None]:
        """
        Stream content from a blob asynchronously.

        Args:
            file_blob_url (str): URL of the blob to stream.

        Returns:
            AsyncGenerator[bytes, None]: Yields chunks of file content asynchronously.

        Raises:
            ValueError: If the provided URL is invalid or missing required parts.
            ResourceNotFoundError: If the blob does not exist.
            Exception: For other unexpected errors.
        """
        if not file_blob_url:
            raise ValueError("Blob URL cannot be empty.")

        try:
            parsed_url = urlparse(file_blob_url)
            if not parsed_url.path:
                raise ValueError("Invalid blob URL: Missing path.")

            # Defensive initialization
            container_name = None
            blob_name = None

            # Extract the blob name from the URL
            # First try the expected recordings container
            try:
                recordings_container = self.config.azure_storage_recordings_container
            except BLOB_STREAM_ERRORS:
                recordings_container = None

            if recordings_container and recordings_container in parsed_url.path:
                container_name = recordings_container
                blob_name = parsed_url.path.split(recordings_container, 1)[-1].lstrip("/")
            else:
                # For transcription files or other assets that might be in different containers,
                # try to extract container and blob name from the URL path
                path_parts = parsed_url.path.strip('/').split('/')
                if len(path_parts) >= 2:
                    # Assume format: /container_name/blob_name or /container_name/folder/blob_name
                    container_name = path_parts[0]
                    blob_name = '/'.join(path_parts[1:])
                    self.logger.warning(
                        "blob_stream_container_mismatch",
                        container=container_name,
                        expected_container=recordings_container,
                    )
                else:
                    raise ValueError(f"Blob URL path format not recognized: {parsed_url.path}")

            if not container_name or not blob_name:
                raise ValueError(f"Could not extract container/blob from URL: {file_blob_url}")
            
            if not blob_name:
                raise ValueError("Could not extract blob name from URL")
            self.logger.debug("blob_stream_name_extracted", blob_name=blob_name)

            # Create an async blob client for the resolved container/blob
            async_blob_client = BlobClient(
                account_url=self.config.azure_storage_account_url,
                container_name=container_name,
                blob_name=blob_name,
                credential=self.credential,
            )

            # Stream the blob content in chunks
            try:
                async with async_blob_client:
                    downloader = await async_blob_client.download_blob()
                    # The downloader provides an async iterator of chunks
                    async for chunk in downloader.chunks():
                        yield chunk
            except ResourceNotFoundError:
                self.logger.error("blob_stream_not_found", blob_name=blob_name)
                raise
            except BLOB_STREAM_ERRORS as e:
                self.logger.error(
                    "blob_stream_failed",
                    blob_name=blob_name,
                    error=str(e),
                    error_type=type(e).__name__,
                    exc_info=True,
                )
                raise

        except ValueError as ve:
            self.logger.warning("blob_stream_validation_failed", error=str(ve))
            raise
        except ResourceNotFoundError as rnfe:
            self.logger.error("blob_stream_not_found", error=str(rnfe))
            raise
        except BLOB_STREAM_ERRORS as e:
            self.logger.error(
                "blob_stream_unexpected_error",
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            raise
