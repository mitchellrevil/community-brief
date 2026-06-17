import io
import os
import tempfile
import pytest
from unittest.mock import Mock

from services.file_processing_service import FileProcessingService


@pytest.mark.unit
class TestFileProcessingServiceDocxExtraction:
    def test_extract_docx_text_with_python_docx(self, file_processing_service):
        # Build an in-memory docx with python-docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        result = file_processing_service._extract_docx_text(doc_bytes)
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_extract_docx_text_fallback_to_docx2txt_when_python_docx_fails(self, monkeypatch, file_processing_service):
        # Simulate python-docx failing and docx2txt returning expected text
        def fake_document(_):
            raise RuntimeError("simulated python-docx failure")

        monkeypatch.setattr('services.file_processing_service.Document', fake_document)
        monkeypatch.setattr('services.file_processing_service.docx2txt', Mock(process=Mock(return_value="fallback text")))

        # Any bytes will do since we're mocking the processing
        result = file_processing_service._extract_docx_text(b"fake-docx-bytes")
        assert result == "fallback text"

    def test_extract_docx_text_handles_unlink_errors(self, monkeypatch, file_processing_service):
        # Simulate python-docx failing, docx2txt succeeds, but os.unlink raises PermissionError
        monkeypatch.setattr('services.file_processing_service.Document', lambda _: (_ for _ in ()).throw(RuntimeError("simulated python-docx failure")))
        monkeypatch.setattr('services.file_processing_service.docx2txt', Mock(process=Mock(return_value="fallback text")))

        def fake_unlink(path):
            raise PermissionError("simulated unlink failure")
        # Patch only in module namespace so we can simulate unlink raising
        monkeypatch.setattr('services.file_processing_service.os.unlink', fake_unlink)

        result = file_processing_service._extract_docx_text(b"fake-docx-bytes")
        assert result == "fallback text"
        # Restore unlink for cleanliness (monkeypatch fixture handles it)
