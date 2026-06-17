import os
from typing import Optional, Any
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from azure.core.exceptions import AzureError
from datetime import UTC, datetime, timedelta
import time
from urllib.parse import urlparse

import structlog

from config import AppConfig

logger = structlog.get_logger(__name__)

STORAGE_SERVICE_ERRORS = (AzureError, RuntimeError, OSError, ValueError, TypeError)
MARKDOWN_PARSE_ERRORS = (ValueError, TypeError, AttributeError)
MARKDOWN_IMPORT_ERRORS = (ImportError, ModuleNotFoundError)


def _preview(text: Optional[str], n: int = 120) -> str:
    if not text:
        return text or ""
    normalized = " ".join(text.replace("\r", " ").replace("\n", " ").split())
    return normalized if len(normalized) <= n else normalized[:n] + "â€¦"


class StorageServiceError(Exception):
    """Custom exception for storage service errors."""
    pass

class StorageService:
    def __init__(self, config: AppConfig, credential: Any = None, blob_service_client: BlobServiceClient = None) -> None:
        """Initialize the StorageService with config, optional credential, and blob service client."""
        self.config = config
        # Lazy import of DefaultAzureCredential to avoid import-time failures
        if credential is not None:
            self.credential = credential
        else:
            try:
                from azure.identity import DefaultAzureCredential
                self.credential = DefaultAzureCredential()
            except (ImportError, ModuleNotFoundError):
                self.credential = None

        # Initialize blob service client
        self.blob_service_client = blob_service_client if blob_service_client is not None else BlobServiceClient(
            account_url=self.config.storage_account_url,
            credential=self.credential,
        )

    def upload_file(self, file_path: str, original_filename: str) -> str:
        """Upload a file to blob storage and return the blob URL."""
        try:
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )            # Generate blob name with date and nested structure including timestamp for uniqueness
            now = datetime.now(UTC)
            current_date = now.strftime("%Y-%m-%d")
            timestamp = now.strftime("%H%M%S_%f")[:-3]  # HHMMSS_milliseconds
            file_name_without_ext = os.path.splitext(original_filename)[0]
            # Include timestamp in both folder and filename to ensure uniqueness
            blob_name = f"{current_date}/{file_name_without_ext}_{timestamp}/{original_filename}"

            blob_client = container_client.get_blob_client(blob_name)

            # Upload the file
            logger.info("storage.file_upload_started", blob_name=blob_name)
            with open(file_path, "rb") as f:
                file_data = f.read()
            blob_client.upload_blob(file_data, overwrite=True)

            return blob_client.url

        except AzureError as e:
            logger.error("storage.file_upload_azure_failed", error=str(e), error_type=type(e).__name__)
            raise StorageServiceError(f"Azure storage error: {str(e)}") from e
        except STORAGE_SERVICE_ERRORS as e:
            logger.error("storage.file_upload_failed", error=str(e), error_type=type(e).__name__)
            raise StorageServiceError(f"Error uploading file: {str(e)}") from e

    def upload_text(
        self, container_name: str, blob_name: str, text_content: str
    ) -> str:
        """Upload text content to blob storage and return the blob URL."""
        try:
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            blob_client = container_client.get_blob_client(blob_name)

            blob_client.upload_blob(text_content.encode("utf-8"), overwrite=True)
            return blob_client.url
        except STORAGE_SERVICE_ERRORS as e:
            logger.error("storage.text_upload_failed", error=str(e), error_type=type(e).__name__)
            raise StorageServiceError(f"Error uploading text: {str(e)}") from e

    def download_text_from_blob(self, blob_url: str) -> str:
        """Download the blob at `blob_url` and return the decoded text."""
        return self._download_blob(blob_url, as_text=True)

    def generate_and_upload_pdf(self, analysis_text: str, blob_url: str) -> str:
        """Generate a PDF from analysis text and upload to blob storage. Return the blob URL."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            import io

            # Create PDF in memory
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)

            # Add content to PDF
            y = 750  # Starting y position
            for line in analysis_text.split("\n"):
                if y < 50:  # Start new page if near bottom
                    c.showPage()
                    y = 750
                c.drawString(50, y, line)
                y -= 15

            c.save()
            pdf_content = buffer.getvalue()

            # Upload PDF
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_url)

            blob_client.upload_blob(pdf_content, overwrite=True)
            return blob_client.url

        except STORAGE_SERVICE_ERRORS as e:
            logger.error("storage.pdf_upload_failed", error=str(e), error_type=type(e).__name__)
            raise StorageServiceError(f"Error generating/uploading PDF: {str(e)}") from e

    def generate_and_upload_docx(self, analysis_text: str, blob_url: str) -> str:
        """Generate a DOCX from analysis text and upload to blob storage. Return the blob URL."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import io

            # Create DOCX in memory
            doc = Document()

            # Parse markdown into tokens (AST-like structure)
            try:
                from markdown_it import MarkdownIt
                import importlib

                # Prefer the commonmark preset and enable linkify where supported
                try:
                    md = MarkdownIt("commonmark", {"linkify": True})
                except MARKDOWN_IMPORT_ERRORS:
                    md = MarkdownIt({"linkify": True})

                # Enable common optional parsing rules when present
                for rule in ("table", "strikethrough"):
                    try:
                        md.enable(rule)
                    except MARKDOWN_PARSE_ERRORS:
                        logger.debug("storage.docx_markdown_rule_unavailable", rule=rule)

                # Attempt to load a task-list plugin from common plugin packages
                tasklist_loaded = False
                tasklist_candidates = [
                    ("mdit_py_plugins.tasklist", "tasklist_plugin"),
                    ("mdit_py_plugins.tasklists", "plugin"),
                    ("mdit_py_plugins.tasklists", "tasklist_plugin"),
                ]
                for mod_name, attr in tasklist_candidates:
                    try:
                        mod = importlib.import_module(mod_name)
                        plugin = getattr(mod, attr)
                        md.use(plugin)
                        tasklist_loaded = True
                        logger.info("storage.docx_tasklist_plugin_loaded", module=mod_name, attribute=attr)
                        break
                    except MARKDOWN_IMPORT_ERRORS:
                        continue
                    except MARKDOWN_PARSE_ERRORS:
                        continue

                if not tasklist_loaded:
                    logger.debug("No tasklist plugin available; task list checkboxes will be handled as plain list items")

                tokens = md.parse(analysis_text)

            except ModuleNotFoundError:
                # If markdown-it-py isn't available in the runtime, fall back to a
                # simple paragraph-based approach so DOCX generation still succeeds
                logger.warning("markdown-it-py not available; falling back to simple paragraph parsing")

                # Render paragraphs directly and upload without trying to parse tokens
                paragraphs = [p for p in analysis_text.split("\n\n") if p.strip()]
                for para in paragraphs:
                    p = doc.add_paragraph()
                    # Reuse the inline renderer by creating a minimal inline token wrapper
                    fake_inline = type("InlineToken", (), {"children": [type("Child", (), {"type": "text", "content": para})]})
                    self._render_inline_tokens(p, fake_inline)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                docx_content = buffer.getvalue()

                logger.info("storage.docx_generated", mode="fallback", size_bytes=len(docx_content))

                # Upload DOCX
                container_client = self.blob_service_client.get_container_client(
                    self.config.storage_recordings_container
                )
                blob_client = container_client.get_blob_client(blob_url)

                blob_client.upload_blob(docx_content, overwrite=True)
                logger.info("storage.docx_upload_completed", blob_url=blob_url)
                return blob_client.url
            except MARKDOWN_PARSE_ERRORS as parse_err:
                logger.error("Failed to parse markdown with markdown-it-py", exc_info=True)
                raise StorageServiceError(f"Markdown parsing error: {parse_err}") from parse_err

            # Track list nesting for style selection
            list_stack: list[dict] = []  # each: {type: 'ordered'|'bullet', level: int}

            i = 0
            while i < len(tokens):
                tok = tokens[i]

                # Headings
                if tok.type == 'heading_open':
                    # next token should be inline with content, followed by heading_close
                    level = 1
                    try:
                        tag = tok.tag  # e.g., 'h1'
                        if tag and tag.startswith('h'):
                            level = max(1, min(3, int(tag[1:])))
                    except MARKDOWN_PARSE_ERRORS:
                        level = 1

                    text = ''
                    if i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                        text = self._collect_plain_text(tokens[i+1])

                    doc.add_heading(text.strip(), level=level)
                    # skip inline and heading_close
                    i += 3
                    continue

                # Paragraphs
                if tok.type == 'paragraph_open':
                    p = doc.add_paragraph()
                    if i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                        self._render_inline_tokens(p, tokens[i+1])
                    i += 3  # paragraph_open, inline, paragraph_close
                    continue

                # Lists
                if tok.type in ('ordered_list_open', 'bullet_list_open'):
                    list_type = 'ordered' if tok.type == 'ordered_list_open' else 'bullet'
                    level = len(list_stack) + 1
                    list_stack.append({'type': list_type, 'level': level})
                    i += 1
                    continue

                if tok.type == 'list_item_open':
                    current = list_stack[-1] if list_stack else {'type': 'bullet', 'level': 1}
                    base_style = 'List Number' if current['type'] == 'ordered' else 'List Bullet'
                    style = base_style if current['level'] == 1 else f"{base_style} {current['level']}"

                    task_state = None
                    if isinstance(getattr(tok, "meta", None), dict) and "checked" in tok.meta:
                        task_state = bool(tok.meta.get("checked"))

                    # Safely resolve the style name: if the specific nested style doesn't exist,
                    # fall back to the base style, otherwise use no style.
                    try:
                        doc_style_names = {s.name for s in doc.styles}
                    except MARKDOWN_PARSE_ERRORS:
                        doc_style_names = set()

                    if style in doc_style_names:
                        use_style = style
                    elif base_style in doc_style_names:
                        use_style = base_style
                    else:
                        use_style = None

                    # Create paragraph with the resolved style (or no style if none found)
                    if use_style:
                        try:
                            p = doc.add_paragraph(style=use_style)
                        except MARKDOWN_PARSE_ERRORS:
                            # Something unexpected about styles; create without style
                            p = doc.add_paragraph()
                    else:
                        p = doc.add_paragraph()

                    if task_state is not None:
                        prefix = "[x] " if task_state else "[ ] "
                        p.add_run(prefix)
                        setattr(p, "_task_prefix_added", True)

                    if i + 1 < len(tokens) and tokens[i+1].type == 'paragraph_open':
                        # inline expected at i+2
                        if i + 2 < len(tokens) and tokens[i+2].type == 'inline':
                            self._render_inline_tokens(p, tokens[i+2])
                        i += 4  # li_open, para_open, inline, para_close
                    elif i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                        self._render_inline_tokens(p, tokens[i+1])
                        i += 2  # li_open, inline
                    else:
                        i += 1
                    continue

                if tok.type in ('ordered_list_close', 'bullet_list_close'):
                    if list_stack:
                        list_stack.pop()
                    i += 1
                    continue

                # Fenced code blocks
                if tok.type in ('fence', 'code_block'):
                    p = doc.add_paragraph()
                    self._render_code_block(p, tok.content or '')
                    i += 1
                    continue

                # Tables
                if tok.type == 'table_open':
                    table_data = self._parse_table_tokens(tokens, i)
                    if table_data:
                        self._render_table(doc, table_data)
                        # Skip all table-related tokens
                        i = table_data['end_index']
                    else:
                        i += 1
                    continue

                # Skip closes and other container tokens
                i += 1

            # Save DOCX to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_content = buffer.getvalue()
            
            logger.info("storage.docx_generated", mode="markdown", size_bytes=len(docx_content))

            # Upload DOCX
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_url)

            blob_client.upload_blob(docx_content, overwrite=True)
            logger.info("storage.docx_upload_completed", blob_url=blob_url)
            return blob_client.url

        except STORAGE_SERVICE_ERRORS as e:
            logger.error(
                "storage.docx_upload_failed",
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            # Log the first 500 chars of analysis text for debugging
            if analysis_text:
                logger.error("storage.docx_analysis_preview", analysis_preview=_preview(analysis_text, n=500))
            raise StorageServiceError(f"Error generating/uploading DOCX: {str(e)}") from e

    def _collect_plain_text(self, inline_token) -> str:
        """Collect plain text from an inline token (concatenate text children)."""
        text_parts = []
        for child in getattr(inline_token, 'children', []) or []:
            if child.type == 'text' or child.type == 'code_inline':
                text_parts.append(child.content)
            elif child.type in ('softbreak', 'hardbreak'):
                text_parts.append('\n')
            elif child.type == 'link_open':
                # next child may contain link text; handled by text tokens
                continue
        return ''.join(text_parts)

    def _render_inline_tokens(self, paragraph, inline_token) -> None:
        """Render inline tokens (**bold**, *italic*, `code`, links, text) into a python-docx paragraph."""
        from docx.shared import Pt
        from docx.shared import RGBColor

        children = getattr(inline_token, 'children', []) or []
        i = 0
        while i < len(children):
            child = children[i]
            
            if child.type == 'text':
                paragraph.add_run(child.content)
                i += 1
            elif child.type == 'code_inline':
                run = paragraph.add_run(child.content)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
            elif child.type == 'strong_open':
                # Gather text until strong_close and skip those tokens
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'strong_close')
                run = paragraph.add_run(text)
                run.bold = True
                i += tokens_consumed
            elif child.type == 'em_open':
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'em_close')
                run = paragraph.add_run(text)
                run.italic = True
                i += tokens_consumed
            elif child.type in ('softbreak', 'hardbreak'):
                paragraph.add_run('\n')
                i += 1
            elif child.type == 'link_open':
                # render link text; ignore URL styling for now
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 'link_close')
                href = self._get_token_attr(child, "href")
                if href:
                    self._add_hyperlink(paragraph, href, text or href)
                else:
                    paragraph.add_run(text)
                i += tokens_consumed
            elif child.type == 's_open':
                # Handle strikethrough (GFM feature)
                text, tokens_consumed = self._gather_text_between_with_count(children, i, 's_close')
                run = paragraph.add_run(text)
                run.font.strike = True
                i += tokens_consumed
            elif child.type == 'task_list_item_checkbox':
                checked = False
                if isinstance(getattr(child, "meta", None), dict):
                    checked = bool(child.meta.get("checked"))
                if not getattr(paragraph, "_task_prefix_added", False):
                    prefix = "[x] " if checked else "[ ] "
                    paragraph.add_run(prefix)
                    setattr(paragraph, "_task_prefix_added", True)
                i += 1
            else:
                # Skip any other token types (like close tags we've already handled)
                i += 1

    def _get_token_attr(self, token, attr_name: str) -> Optional[str]:
        """Return a token attribute value if present."""
        if hasattr(token, "attrGet"):
            try:
                value = token.attrGet(attr_name)
                if value:
                    return value
            except MARKDOWN_PARSE_ERRORS:
                pass

        attrs = getattr(token, "attrs", None) or []
        for key, value in attrs:
            if key == attr_name:
                return value
        return None

    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        """Add a hyperlink to a paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        new_run.append(r_pr)

        new_text = OxmlElement("w:t")
        new_text.text = text
        new_run.append(new_text)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _gather_text_between(self, children: list, start_token, end_type: str) -> str:
        """Helper to gather plain text between start_token and the corresponding end_type token."""
        collecting = False
        text_parts = []
        open_count = 0
        for t in children:
            if t is start_token:
                collecting = True
                open_count = 1
                continue
            if collecting:
                if t.type == start_token.type:
                    open_count += 1
                elif t.type == end_type:
                    open_count -= 1
                    if open_count == 0:
                        break
                elif t.type in ('text', 'code_inline'):
                    text_parts.append(t.content)
        return ''.join(text_parts)
    
    def _gather_text_between_with_count(self, children: list, start_index: int, end_type: str) -> tuple[str, int]:
        """
        Helper to gather plain text between start_token and the corresponding end_type token.
        Returns tuple of (text, number of tokens consumed including open and close tags).
        """
        text_parts = []
        open_count = 1  # We've already seen the opening tag
        start_token_type = children[start_index].type
        tokens_consumed = 1  # Count the opening tag
        
        for i in range(start_index + 1, len(children)):
            t = children[i]
            tokens_consumed += 1
            
            if t.type == start_token_type:
                # Nested opening tag of same type
                open_count += 1
            elif t.type == end_type:
                open_count -= 1
                if open_count == 0:
                    # Found the matching closing tag
                    break
            elif t.type in ('text', 'code_inline'):
                text_parts.append(t.content)
        
        return ''.join(text_parts), tokens_consumed

    def _render_code_block(self, paragraph, content: str) -> None:
        """Render fenced/code block content in monospace formatting."""
        from docx.shared import Pt
        run = paragraph.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    def _parse_table_tokens(self, tokens: list, start_index: int) -> Optional[dict]:
        """Parse table tokens and return table data structure."""
        rows = []
        current_row = []
        in_header = True
        i = start_index
        
        while i < len(tokens):
            tok = tokens[i]
            
            if tok.type == 'table_open':
                i += 1
                continue
                
            if tok.type == 'thead_open':
                in_header = True
                i += 1
                continue
                
            if tok.type == 'tbody_open':
                in_header = False
                i += 1
                continue
                
            if tok.type == 'tr_open':
                current_row = []
                i += 1
                continue
                
            if tok.type == 'th_open' or tok.type == 'td_open':
                # Collect inline content for cell
                cell_content = ''
                if i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                    cell_content = self._collect_plain_text(tokens[i+1])
                current_row.append(cell_content)
                # Skip th/td_open, inline, th/td_close
                i += 3
                continue
                
            if tok.type == 'tr_close':
                if current_row:
                    rows.append({
                        'cells': current_row,
                        'is_header': in_header
                    })
                i += 1
                continue
                
            if tok.type == 'thead_close' or tok.type == 'tbody_close':
                i += 1
                continue
                
            if tok.type == 'table_close':
                return {
                    'rows': rows,
                    'end_index': i + 1
                }
                
            i += 1
            
        return None

    def _render_table(self, doc, table_data: dict) -> None:
        """Render a table in the Word document."""
        from docx.shared import Inches
        
        rows = table_data['rows']
        if not rows:
            return
            
        # Create table with correct number of rows and columns
        num_rows = len(rows)
        num_cols = max(len(row['cells']) for row in rows) if rows else 0
        
        if num_rows == 0 or num_cols == 0:
            return
            
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Set column widths (distribute evenly)
        for col in table.columns:
            col.width = Inches(1.0)
        
        # Fill table data
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_content in enumerate(row_data['cells']):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_content
                    
                    # Make header rows bold
                    if row_data['is_header']:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _parse_blob_url(self, blob_url: str) -> tuple[str, str]:
        """Extract the container and blob name from a blob URL.

        Raises ValueError for malformed URLs.
        """
        parsed = urlparse(blob_url)

        # Path is /{container}/{blob_path}
        path_parts = parsed.path.lstrip("/").split("/", 1)
        if len(path_parts) < 2:
            raise ValueError("Invalid blob url: missing container or blob name")
        container = path_parts[0]
        blob_name = path_parts[1]

        # Remove any SAS token from blob_name
        if "?" in blob_name:
            blob_name = blob_name.split("?", 1)[0]

        return container, blob_name

    def generate_sas_url(self, blob_url: str, expiry_hours: int = 1) -> str:
        """Generate a read-only SAS URL for a blob.

        This will use a user delegation key if the BlobServiceClient has
        a token credential (DefaultAzureCredential), otherwise it will
        attempt to use the account key from the environment.
        """
        try:
            container, blob_name = self._parse_blob_url(blob_url)

            # Expiry time for SAS
            expiry = datetime.now(UTC) + timedelta(hours=expiry_hours)

            # Account name from the configured account URL
            account_name = None
            if self.config.storage_account_url:
                # storage_account_url like https://<account>.blob.core.windows.net
                parsed_account = urlparse(self.config.storage_account_url)
                account_name = parsed_account.netloc.split(".")[0]

            # Prefer user delegation key (when using DefaultAzureCredential)
            if self.credential is not None and hasattr(self.blob_service_client, "get_user_delegation_key"):
                start = time.perf_counter()
                # Use user delegation key; start now and expire at expiry
                key = self.blob_service_client.get_user_delegation_key(datetime.now(UTC), expiry)
                sas = generate_blob_sas(
                    account_name=account_name,
                    container_name=container,
                    blob_name=blob_name,
                    user_delegation_key=key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                )
                end = time.perf_counter()
                logger.info(
                    "storage.sas_generated",
                    blob_url=blob_url,
                    elapsed_ms=int((end - start) * 1000),
                    method="user_delegation_key",
                )
            else:
                start = time.perf_counter()
                account_key = self.config.storage_account_key
                sas = generate_blob_sas(
                    account_name=account_name,
                    container_name=container,
                    blob_name=blob_name,
                    account_key=account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                )
                end = time.perf_counter()
                logger.info(
                    "storage.sas_generated",
                    blob_url=blob_url,
                    elapsed_ms=int((end - start) * 1000),
                    method="account_key",
                )

            if not sas:
                raise StorageServiceError("Failed to generate SAS token")

            # Append SAS to original blob URL (strip existing query if present)
            delimiter = "&" if "?" in blob_url else "?"
            return f"{blob_url}{delimiter}{sas}"

        except STORAGE_SERVICE_ERRORS as e:
            logger.error("Failed to generate SAS URL", exc_info=True)
            raise StorageServiceError(f"Failed to generate SAS URL: {str(e)}") from e

    def _download_blob(self, blob_url: str, as_text: bool = False) -> Optional[Any]:
        """Internal helper to download blob content synchronously.

        Returns decoded text when as_text=True, raw bytes when as_text=False,
        or None on failure.
        """
        if not blob_url:
            return None
        try:
            container, blob_name = self._parse_blob_url(blob_url)
            container_client = self.blob_service_client.get_container_client(container)
            blob_client = container_client.get_blob_client(blob_name)

            stream = blob_client.download_blob()
            blob_data = stream.readall()

            if as_text:
                try:
                    return blob_data.decode("utf-8")
                except UnicodeDecodeError:
                    # Fallback to latin-1 if utf-8 fails
                    try:
                        return blob_data.decode("latin-1")
                    except UnicodeDecodeError:
                        logger.exception("Failed to decode blob as text")
                        return None

            return blob_data
        except STORAGE_SERVICE_ERRORS as exc:
            logger.error("Failed to download blob content", exc_info=True)
            return None

